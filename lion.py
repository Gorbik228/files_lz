import docx
import matplotlib.pyplot as plt
from collections import Counter

lion = docx.Document('lion.docx') #делаем произведение
doc = docx.Document() #файл для вывода таблицы
text = [] #лист, чтобы добавить в него текст
letters = ['а', 'б', 'в', 'г', 'д', 'е', 'ё', 'ж', 'з', 'и', 'й', 'к', 'л', 'м', 'н', 'о', 'п', 'р', 'с', 'т', 'у', 'ф', 'х', 'ц',
            'ч', 'ш', 'щ', 'ь', 'ы', 'ъ', 'э', 'ю', 'я', 'q', 'w', 'e', 'r', 't', 'y', 'u', 'i', 'o', 'p', 'a', 's', 'd', 'f', 'g', 'h',
            'j', 'k', 'l', 'z', 'x', 'c', 'v', 'b', 'n', 'm', '1', '2', '3', '4', '5', '6', '7', '8', '9', '0']
#все элементы, которые могут входить
table = doc.add_table(rows = 1, cols = 3) #создаем таблицу 
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Слово' #подписываем 
hdr_cells[1].text = 'Частота встречи, раз'
hdr_cells[2].text = 'Частота встречи, %'

for paragraph in lion.paragraphs:
    text.append(paragraph.text) #добавлякм весь наш текст в созданную переменную

text = '\n'.join(text) #делаем весь текст в одну строку
text_word = text.lower() #в новую переменную добавляем текст только маленькими буквами
text_word = text_word.split() #разбираем его на слова

filtered_words = [
    word.strip(" ,.!?()[]{}'\"")  #удаляем пробелы и знаки препинания
    for word in text_word #проходим по всем словам из списка
    if len(word) > 0 and all(letter in letters for letter in word) #проверяем длину слова и все буквы, чтобы они входили в список 
]

#считаем частоту уникальных слов
word_counts = Counter(filtered_words)

#создаем переменную с количеством слов
total_words = len(filtered_words)

for word, count in word_counts.items(): #проходим по всем значениям 
    percentage = (count / total_words) * 100 #создаем переменную, которая будет отображать частоту в процентах
    new_row = table.add_row() #создаем новую строку
    new_row.cells[0].text = word #заполняем первую ячейку
    new_row.cells[1].text = str(count) #заполняем вторую ячейку
    new_row.cells[2].text = str(percentage) #заполняем третью ячейку

doc.save('table.docx') #сохраняем файл с таблицей

#фильтруем текст, оставляем только маленькие русские буквы
filtered_text = ''.join(filter(lambda x: 'а' <= x <= 'я', text.lower()))

#считаем количество каждой буквы
letter_counts = Counter(filtered_text)

#сортируем буквы по алфавиту
sorted_letters = sorted(letter_counts.keys())
sorted_counts = [letter_counts[letter] for letter in sorted_letters]

#создание графика
plt.figure(figsize=(12, 6))
plt.bar(sorted_letters, sorted_counts, color='skyblue')
plt.xlabel('Буквы')
plt.ylabel('Количество')
plt.title('Количество строчных русских букв в тексте')
plt.grid(axis='y')

# Отображаем график
plt.tight_layout()
plt.show()